import os
import logging

logger = logging.getLogger(__name__)

# Try loading heavy parsing packages, log warning if they are not installed.
try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False
    logger.warning("pdfplumber not installed. PDF parsing will fallback to PyPDF2.")

try:
    from pypdf import PdfReader
    PYPDF_AVAILABLE = True
except ImportError:
    PYPDF_AVAILABLE = False
    logger.warning("pypdf not installed. PDF parsing will fail.")

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not installed. DOCX parsing will fail.")

class ResumeParser:
    """Reads PDF, DOCX, and TXT files and extracts their raw string content."""
    def __init__(self):
        pass

    def parse(self, filepath):
        """Route parsing request based on file extension."""
        if not filepath or not os.path.exists(filepath):
            raise FileNotFoundError(f"File not found: {filepath}")
            
        ext = filepath.rsplit('.', 1)[1].lower() if '.' in filepath else ''
        
        raw_text = ""
        if ext == 'pdf':
            raw_text = self.parse_pdf(filepath)
        elif ext == 'docx':
            raw_text = self.parse_docx(filepath)
        elif ext == 'txt':
            raw_text = self.parse_txt(filepath)
        else:
            raise ValueError(f"Unsupported file format: {ext}")
            
        file_size = os.path.getsize(filepath)
        
        return {
            'raw_text': raw_text,
            'file_type': ext,
            'file_size': file_size,
            'filename': os.path.basename(filepath)
        }

    def parse_pdf(self, filepath):
        """Extract text from PDF, trying pdfplumber first, then PyPDF2."""
        text = ""
        
        # Try pdfplumber
        if PDFPLUMBER_AVAILABLE:
            try:
                with pdfplumber.open(filepath) as pdf:
                    for page in pdf.pages:
                        extracted = page.extract_text()
                        if extracted:
                            text += extracted + "\n"
                if text.strip():
                    return text.strip()
            except Exception as e:
                logger.error(f"pdfplumber failed for {filepath}: {e}")
                
        # Try pypdf fallback
        if PYPDF_AVAILABLE:
            try:
                reader = PdfReader(filepath)
                text = ""
                for page in reader.pages:
                    extracted = page.extract_text()
                    if extracted:
                        text += extracted + "\n"
                return text.strip()
            except Exception as e:
                logger.error(f"pypdf failed for {filepath}: {e}")
                
        raise RuntimeError("No working PDF parsing engine available or file is corrupted.")

    def parse_docx(self, filepath):
        """Extract text paragraphs from a Word document (.docx)."""
        if not DOCX_AVAILABLE:
            raise RuntimeError("python-docx package is required for DOCX parsing.")
            
        try:
            doc = docx.Document(filepath)
            paragraphs = []
            
            # Read paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text.strip())
                    
            # Read table cells if present (many candidates format resumes in tables)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraphs.append(cell.text.strip())
                            
            return "\n".join(paragraphs)
        except Exception as e:
            logger.error(f"python-docx failed for {filepath}: {e}")
            raise RuntimeError(f"Failed parsing DOCX: {e}")

    def parse_txt(self, filepath):
        """Read standard UTF-8 or ANSI text files."""
        try:
            with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read().strip()
        except Exception as e:
            logger.error(f"TXT read failed for {filepath}: {e}")
            raise RuntimeError(f"Failed parsing TXT: {e}")
